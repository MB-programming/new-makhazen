#!/usr/bin/env python3
"""
DOCX → HTML converter
Usage: python3 convert_docx.py <file.docx>
Outputs JSON: {"title": "...", "excerpt": "...", "body": "<html>..."}
"""

import sys
import json
import re
from docx import Document
from docx.oxml.ns import qn

def get_paragraph_html(para):
    style = para.style.name if para.style else ''
    text  = para.text.strip()

    if not text:
        return ''

    # Build inline HTML with run formatting
    html_runs = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        t = t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        if run.bold:
            t = f'<strong>{t}</strong>'
        if run.italic:
            t = f'<em>{t}</em>'
        if run.underline:
            t = f'<u>{t}</u>'
        html_runs.append(t)

    inner = ''.join(html_runs) if html_runs else text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    if re.match(r'Heading 1', style, re.I) or style == 'Title':
        return f'<h2>{inner}</h2>'
    elif re.match(r'Heading 2', style, re.I):
        return f'<h2>{inner}</h2>'
    elif re.match(r'Heading 3', style, re.I):
        return f'<h3>{inner}</h3>'
    elif re.match(r'List', style, re.I):
        return f'<li>{inner}</li>'
    else:
        return f'<p>{inner}</p>'

def convert(path):
    doc = Document(path)
    parts = []
    in_list = False

    for para in doc.paragraphs:
        style = para.style.name if para.style else ''
        text  = para.text.strip()

        is_list = re.match(r'List', style, re.I)

        if is_list and not in_list:
            parts.append('<ul>')
            in_list = True
        elif not is_list and in_list:
            parts.append('</ul>')
            in_list = False

        html = get_paragraph_html(para)
        if html:
            parts.append(html)

    if in_list:
        parts.append('</ul>')

    body = '\n'.join(parts)

    # Extract title = first non-empty paragraph
    title = ''
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            title = t
            break

    # Excerpt = second non-empty paragraph (plain text)
    excerpt = ''
    count = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            count += 1
            if count == 2:
                excerpt = t
                break

    return {"title": title, "excerpt": excerpt, "body": body}

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file specified"}))
        sys.exit(1)
    try:
        result = convert(sys.argv[1])
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)
